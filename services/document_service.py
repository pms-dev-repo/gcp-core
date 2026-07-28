from __future__ import annotations

import platform
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from core.config import get_active_client_code, load_client_config


BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_ROOT = BASE_DIR / "templates"
GENERATED_ROOT = BASE_DIR / "generated"


def _client_paths() -> tuple[str, Path, Path]:
    client_code = get_active_client_code()
    config = load_client_config(client_code)
    client = config.get("client", {})
    templates_folder = str(client.get("templates_folder") or client_code)
    generated_folder = str(client.get("generated_folder") or client_code)
    return (
        client_code,
        TEMPLATES_ROOT / templates_folder,
        GENERATED_ROOT / generated_folder,
    )


def _replace_tokens_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replace placeholders even when Word splits them across several runs."""
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if not any(token in full_text for token in replacements):
        return

    updated_text = full_text
    for token, value in replacements.items():
        updated_text = updated_text.replace(token, value)

    paragraph.runs[0].text = updated_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_tokens_in_table(table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_tokens_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                _replace_tokens_in_table(nested_table, replacements)


def _replace_tokens_in_headers_and_footers(
    document,
    replacements: dict[str, str],
) -> None:
    for section in document.sections:
        for area in (section.header, section.footer):
            for paragraph in area.paragraphs:
                _replace_tokens_in_paragraph(paragraph, replacements)
            for table in area.tables:
                _replace_tokens_in_table(table, replacements)


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", value.strip())
    return cleaned.strip("_") or "guest"


def _resolve_template(guest: dict[str, Any], document_kind: str | None = None) -> tuple[Path, str]:
    movement = str(guest.get("movement", "")).strip().lower()
    template = guest.get("template", {})
    configured_name = str(template.get("file_name", "")).strip()

    kind = (document_kind or "").strip().lower()
    if kind == "confirmation":
        default_name = "confirmation_standard_en.docx"
        fallback_name = "arrival_standard_en.docx"
        default_prefix = "CONFIRMATION"
    elif kind == "cancellation":
        default_name = "cancellation_standard_en.docx"
        fallback_name = "departure_standard_en.docx"
        default_prefix = "CANCELLATION"
    elif movement == "arrivals":
        default_name = "arrival_standard_en.docx"
        fallback_name = default_name
        default_prefix = "ARRIVAL"
    elif movement == "departures":
        default_name = "departure_standard_en.docx"
        fallback_name = default_name
        default_prefix = "DEPARTURE"
    else:
        raise ValueError(f"Unsupported movement: {guest.get('movement')}")

    _, templates_dir, _ = _client_paths()
    template_name = configured_name or default_name
    template_path = templates_dir / template_name

    if not template_path.exists():
        fallback_path = templates_dir / fallback_name
        if fallback_path.exists():
            template_path = fallback_path
        else:
            raise FileNotFoundError(
                f"Template not found: {template_path}. "
                f"Fallback also not found: {fallback_path}"
            )

    prefix = str(template.get("document_prefix") or default_prefix).upper()
    return template_path, prefix


def generate_guest_document(guest: dict[str, Any], document_kind: str | None = None) -> Path:
    template_path, document_prefix = _resolve_template(guest, document_kind=document_kind)

    stay = guest.get("stay", {})
    transport = guest.get("transport", {})
    letter_date = guest.get("letter_date") or datetime.now().strftime("%d %B %Y")

    replacements = {
        "{{letter_date}}": str(letter_date),
        "{{salutation}}": str(guest.get("salutation") or ""),
        "{{guest_full_name}}": str(guest.get("full_name") or ""),
        "{{guest_first_name}}": str(guest.get("first_name") or ""),
        "{{guest_last_name}}": str(guest.get("last_name") or ""),
        "{{room_number}}": str(guest.get("room") or "To be assigned"),
        "{{arrival_date}}": str(stay.get("arrival_date") or "To be confirmed"),
        "{{departure_date}}": str(stay.get("departure_date") or "To be confirmed"),
        "{{eta}}": str(transport.get("eta") or "To be confirmed"),
        "{{confirmation_number}}": str(
            guest.get("confirmation_number") or ""
        ),
        "{{email}}": str(guest.get("email") or ""),
        "{{cancellation_date}}": str(guest.get("cancellation_date") or ""),
        "{{cancellation_reason}}": str(guest.get("cancellation_reason") or ""),
    }

    document = Document(template_path)

    for paragraph in document.paragraphs:
        _replace_tokens_in_paragraph(paragraph, replacements)
    for table in document.tables:
        _replace_tokens_in_table(table, replacements)
    _replace_tokens_in_headers_and_footers(document, replacements)

    _, _, generated_dir = _client_paths()
    generated_dir.mkdir(parents=True, exist_ok=True)

    safe_name = _safe_filename(str(guest.get("full_name") or "guest"))
    confirmation = _safe_filename(
        str(guest.get("confirmation_number") or "no_confirmation")
    )
    filename = f"{document_prefix}_{confirmation}_{safe_name}.docx"

    output_path = generated_dir / filename
    document.save(output_path)
    return output_path


def read_generated_document(path_value: str | Path) -> bytes:
    path = Path(path_value)
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"Generated document not found: {path}")
    return path.read_bytes()


def _validate_docx_path(docx_path_value: str | Path) -> Path:
    docx_path = Path(docx_path_value).resolve()

    if not docx_path.exists() or not docx_path.is_file():
        raise FileNotFoundError(
            f"Generated Word document not found: {docx_path}"
        )

    if docx_path.suffix.lower() != ".docx":
        raise ValueError("Only DOCX files can be converted to PDF.")

    return docx_path


def _create_word_application():
    """Create one hidden Microsoft Word instance for PDF conversion."""
    if platform.system().lower() != "windows":
        raise RuntimeError("Microsoft Word automation is only available on Windows.")

    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError(
            "PDF generation on Windows requires pywin32. "
            "Install it with: python -m pip install pywin32"
        ) from exc

    pythoncom.CoInitialize()

    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        return word, pythoncom
    except Exception:
        pythoncom.CoUninitialize()
        raise


def _close_word_application(word, pythoncom_module) -> None:
    """Close Word and release the COM apartment safely."""
    try:
        if word is not None:
            word.Quit()
    finally:
        if pythoncom_module is not None:
            pythoncom_module.CoUninitialize()


def _convert_docx_with_open_word(word, docx_path: Path) -> Path:
    """Convert one DOCX using an already-open Microsoft Word instance."""
    pdf_path = docx_path.with_suffix(".pdf")

    if pdf_path.exists():
        pdf_path.unlink()

    document = None

    try:
        document = word.Documents.Open(
            str(docx_path),
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
        )

        # 17 = wdFormatPDF
        document.SaveAs2(str(pdf_path), FileFormat=17)
    except Exception as exc:
        raise OSError(
            "Microsoft Word could not convert the document to PDF. "
            f"Original error: {exc}"
        ) from exc
    finally:
        if document is not None:
            try:
                # 0 = wdDoNotSaveChanges
                document.Close(SaveChanges=0)
            except Exception:
                pass

    if not pdf_path.exists() or not pdf_path.is_file():
        raise OSError("Microsoft Word did not create the PDF file.")

    return pdf_path


def _convert_pdf_with_word(docx_path: Path) -> Path:
    """Convert one DOCX to PDF using a temporary Microsoft Word session."""
    word = None
    pythoncom_module = None

    try:
        word, pythoncom_module = _create_word_application()
        return _convert_docx_with_open_word(word, docx_path)
    finally:
        _close_word_application(word, pythoncom_module)

def _find_libreoffice_executable() -> str:
    for executable_name in ("libreoffice", "soffice"):
        executable_path = shutil.which(executable_name)
        if executable_path:
            return executable_path

    raise RuntimeError(
        "LibreOffice is not installed or is not available in PATH. "
        "On Streamlit Community Cloud, add 'libreoffice' to packages.txt."
    )


def _convert_pdf_with_libreoffice(docx_path: Path) -> Path:
    libreoffice = _find_libreoffice_executable()
    output_dir = docx_path.parent
    pdf_path = docx_path.with_suffix(".pdf")

    if pdf_path.exists():
        pdf_path.unlink()

    with tempfile.TemporaryDirectory(prefix="gcp_lo_profile_") as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()

        command = [
            libreoffice,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ]

        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )

    if result.returncode != 0:
        details = (result.stderr or result.stdout or "Unknown LibreOffice error").strip()
        raise OSError(
            "LibreOffice could not convert the document to PDF. "
            f"Exit code: {result.returncode}. Details: {details}"
        )

    if not pdf_path.exists() or not pdf_path.is_file():
        details = (result.stdout or result.stderr or "No output returned").strip()
        raise OSError(
            "LibreOffice finished without creating the expected PDF file. "
            f"Details: {details}"
        )

    return pdf_path


def generate_guest_pdf(docx_path_value: str | Path) -> Path:
    """
    Windows: Microsoft Word through docx2pdf.
    Linux / Streamlit Community Cloud: LibreOffice headless.
    macOS: Word first, then LibreOffice fallback.
    """
    docx_path = _validate_docx_path(docx_path_value)
    system_name = platform.system().lower()

    if system_name == "windows":
        return _convert_pdf_with_word(docx_path)

    if system_name == "darwin":
        try:
            return _convert_pdf_with_word(docx_path)
        except Exception:
            return _convert_pdf_with_libreoffice(docx_path)

    return _convert_pdf_with_libreoffice(docx_path)


def open_pdf_in_new_tab(pdf_path_value: str | Path) -> None:
    import webbrowser

    pdf_path = Path(pdf_path_value).resolve()

    if not pdf_path.exists() or not pdf_path.is_file():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    opened = webbrowser.open_new_tab(pdf_path.as_uri())
    if not opened:
        raise OSError("The browser could not open the PDF in a new tab.")


def generate_bulk_documents(
    guests: list[dict[str, Any]],
) -> dict[str, Any]:
    results: list[dict[str, Any]] = []
    generated = 0
    failed = 0

    for guest in guests:
        guest_id = guest["id"]
        full_name = guest.get("full_name") or "Unknown guest"

        try:
            output_path = generate_guest_document(guest)
            generated += 1
            results.append(
                {
                    "guest_id": guest_id,
                    "Guest": full_name,
                    "Status": "Generated",
                    "Path": str(output_path),
                    "Message": "Guest letter generated successfully",
                }
            )
        except Exception as exc:
            failed += 1
            results.append(
                {
                    "guest_id": guest_id,
                    "Guest": full_name,
                    "Status": "Failed",
                    "Path": "",
                    "Message": str(exc),
                }
            )

    return {
        "total": len(guests),
        "generated": generated,
        "failed": failed,
        "skipped": 0,
        "results": results,
    }


def generate_bulk_pdfs(
    guests: list[dict[str, Any]],
    generated_documents: dict[str, Any],
) -> dict[str, Any]:
    """
    Convert all selected DOCX files to PDF.

    On Windows, one hidden Microsoft Word instance is reused for the entire
    batch. This avoids repeated COM startup/shutdown failures during bulk
    conversion. Other operating systems continue using LibreOffice.
    """
    results: list[dict[str, Any]] = []
    generated = 0
    failed = 0
    skipped = 0

    pending: list[tuple[dict[str, Any], Path]] = []

    for guest in guests:
        guest_id = guest["id"]
        full_name = guest.get("full_name") or "Unknown guest"
        docx_path_value = generated_documents.get(guest_id)

        if not docx_path_value:
            skipped += 1
            results.append(
                {
                    "guest_id": guest_id,
                    "Guest": full_name,
                    "Status": "Skipped",
                    "Path": "",
                    "Message": "Generate the guest letter first",
                }
            )
            continue

        docx_path = Path(docx_path_value)

        if not docx_path.exists() or not docx_path.is_file():
            skipped += 1
            results.append(
                {
                    "guest_id": guest_id,
                    "Guest": full_name,
                    "Status": "Skipped",
                    "Path": "",
                    "Message": f"DOCX file not found: {docx_path}",
                }
            )
            continue

        pending.append((guest, docx_path))

    system_name = platform.system().lower()

    if system_name == "windows" and pending:
        word = None
        pythoncom_module = None

        try:
            word, pythoncom_module = _create_word_application()

            for guest, docx_path in pending:
                guest_id = guest["id"]
                full_name = guest.get("full_name") or "Unknown guest"

                try:
                    pdf_path = _convert_docx_with_open_word(word, docx_path)
                    generated += 1
                    results.append(
                        {
                            "guest_id": guest_id,
                            "Guest": full_name,
                            "Status": "Generated",
                            "Path": str(pdf_path),
                            "Message": "PDF generated successfully",
                        }
                    )
                except Exception as exc:
                    failed += 1
                    results.append(
                        {
                            "guest_id": guest_id,
                            "Guest": full_name,
                            "Status": "Failed",
                            "Path": "",
                            "Message": str(exc),
                        }
                    )
        except Exception as exc:
            # Word could not start. Mark every pending item as failed.
            for guest, _ in pending:
                failed += 1
                results.append(
                    {
                        "guest_id": guest["id"],
                        "Guest": guest.get("full_name") or "Unknown guest",
                        "Status": "Failed",
                        "Path": "",
                        "Message": f"Microsoft Word could not start: {exc}",
                    }
                )
        finally:
            _close_word_application(word, pythoncom_module)

    else:
        for guest, docx_path in pending:
            guest_id = guest["id"]
            full_name = guest.get("full_name") or "Unknown guest"

            try:
                pdf_path = generate_guest_pdf(docx_path)
                generated += 1
                results.append(
                    {
                        "guest_id": guest_id,
                        "Guest": full_name,
                        "Status": "Generated",
                        "Path": str(pdf_path),
                        "Message": "PDF generated successfully",
                    }
                )
            except Exception as exc:
                failed += 1
                results.append(
                    {
                        "guest_id": guest_id,
                        "Guest": full_name,
                        "Status": "Failed",
                        "Path": "",
                        "Message": str(exc),
                    }
                )

    return {
        "total": len(guests),
        "generated": generated,
        "failed": failed,
        "skipped": skipped,
        "results": results,
    }

